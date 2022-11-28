
#!/usr/bin/env python3
# doc.save('demo.docx')
from docx import Document
from docx.dml.color import ColorFormat
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import re


# document = Document('SpottyProduce.docx')
# doc = Document()

# for i in range(len(document.paragraphs)):
#     for j in range(len(document.paragraphs[i].runs)):
#         try: 
#             if(document.paragraphs[i].runs[j].bold):
#                 #  or document.paragraphs[i].runs[j].underline):
#                 doc.add_paragraph(document.paragraphs[i].runs[j].text)
#             elif(document.paragraphs[i].runs[j].element.rPr.xpath('w:shd')[0].get(qn('w:fill'))):
#                 doc.add_paragraph(document.paragraphs[i].runs[j].text)   
#         except (AttributeError, IndexError):
#             continue

# doc.save('multipleParagraphs.docx')

def find_string(filename):
    counter = 0
    doc = Document(filename)
    for p in doc.paragraphs:
        # if re.match(r'[0-9]') in p.text:
            inline = p.runs
            for i in range(len(inline)):
                if (re.search(r'[0-9]+(.+)(\){1})', inline[i].text) or re.search(r'(Vocabulary)', inline[i].text)):
                    counter += 1
                    print("Found!", inline[i].text, counter)


find_string('SpottyProduce.docx')


# def delete_paragraph(paragraph):
#     p = paragraph._element
#     p.getparent().remove(p)
#     p._p = p._element = None

# def delete_run(run):
#     p = run._element
#     p.getparent().remove(p)
#     p._p = p._element = None

# for i in reversed(range(len(document.paragraphs))):
#     for j in reversed(range(len(document.paragraphs[i].runs))):
#         try:
#             if(document.paragraphs[i].runs[j].text.isnumeric() or document.paragraphs[i].runs[j].font.superscript):
#                 document.paragraphs[i].runs[j].text = ''
#             elif(document.paragraphs[i].runs[j].bold or document.paragraphs[i].runs[j].underline):
#                 continue
#             else:
#                 document.paragraphs[i].runs[j].text = ''
#         except IndexError:
#             continue

# for i in reversed(range(len(document.paragraphs))):
#     if (len(document.paragraphs[i].text.strip()) == 0):
#         delete_paragraph(document.paragraphs[i])

# for i in reversed(range(len(document.paragraphs))):
#     for j in reversed(range(len(document.paragraphs[i].runs))):
#         if (len(document.paragraphs[i].runs[j].text.strip()) == 0):
#             delete_run(document.paragraphs[i].runs[j])

# counter = 1

# for i in range(len(document.paragraphs)):
#     for j in range(len(document.paragraphs[i].runs)):
#         document.paragraphs[i].runs[j].text = str(counter) + ". " + "\n" + document.paragraphs[i].runs[j].text + ": " + "\n"
#         counter += 1

# for i in range(len(document.paragraphs)):
#     for j in range(len(document.paragraphs[i].runs)):
#         if(document.paragraphs[i].text[0].isnumeric()):
#             document.paragraphs[i].runs[j].font.color.rgb = None

# document.save('5_Dog Years_廣.docx')

# doc = Document('5_The Invisible Girl_廣.docx')
# for i in range(len(doc.paragraphs)):
#     for j in range(len(doc.paragraphs[i].runs)):
#         try:
#             if(doc.paragraphs[i].runs[j].element.rPr.xpath('w:shd')[0].get(qn('w:fill')) != None):
#                 print(doc.paragraphs[i].runs[j].text)   
#         except IndexError:
#             continue
# print(doc.paragraphs[28].runs[4].element.rPr.xpath('w:shd')[0].get(qn('w:fill')) == "FFFF00")
# print(doc.paragraphs[95].runs[2].text)
